import argparse
import json
import re
from pathlib import Path


def extract_paper_title(file_path: str) -> str:
    """
    从论文文件中提取标题
    支持PDF、TXT、DOCX等格式的论文标题提取
    """
    path = Path(file_path)
    file_extension = path.suffix.lower()

    try:
        if file_extension == '.pdf':
            return _extract_title_from_pdf(file_path)
        elif file_extension == '.txt':
            return _extract_title_from_txt(file_path)
        elif file_extension == '.docx':
            return _extract_title_from_docx(file_path)
        else:
            # 对于不支持的格式，返回文件名作为标题
            return path.stem
    except Exception as e:
        print(f"提取论文标题失败 ({file_path}): {e}")
        return path.stem


def _extract_title_from_txt(file_path: str) -> str:
    """从TXT文件中提取标题"""
    try:
        with open(file_path, encoding='utf-8') as f:
            content = f.read()

        # 尝试匹配常见的标题格式
        patterns = [
            r'^[A-Z][A-Za-z\s,\-:;]+$',  # 首行大写
            r'^[0-9]+\.\s+[A-Z][A-Za-z\s,\-:;]+$',  # 数字开头
            r'^Title:\s*(.+)',  # Title: 开头
            r'^第[一二三四五六七八九十\d]+章\s*[：:]\s*(.+)',  # 中文章节标题
        ]

        lines = content.split('\n')
        for line in lines[:10]:  # 只检查前10行
            line = line.strip()
            if not line:
                continue

            for pattern in patterns:
                match = re.match(pattern, line)
                if match:
                    return match.group(1).strip()

        # 如果没有找到匹配的标题，返回第一行非空内容
        for line in lines[:10]:
            line = line.strip()
            if line and len(line) > 5:  # 标题通常较长
                return line

        return Path(file_path).stem

    except Exception as e:
        print(f"从TXT文件提取标题失败: {e}")
        return Path(file_path).stem


def _is_likely_not_title(line: str) -> bool:
    """判断是否明显不是标题"""
    # 纯数字或页码
    if re.match(r'^\d+$', line):
        return True
    # 日期格式
    if re.match(r'^\d{4}[-/]\d{1,2}[-/]\d{1,2}', line):
        return True
    # 页眉页脚常见模式
    if re.match(r'^(page|vol|volume|issue|no\.|number)', line, re.IGNORECASE):
        return True
    # 作者信息（含 @ 或邮箱）
    if '@' in line or re.search(r'\.edu|\.com|\.org', line):
        return True
    # 以句号结尾（通常是摘要或正文）
    if line.endswith('.') and len(line) > 30:
        return True
    # 参考文献格式
    if re.match(r'^\[\d+\]', line):
        return True
    return False


def _title_score(line: str) -> int:
    """计算标题得分，越高越可能是标题"""
    score = 0

    # 长度适中（10-150字符）
    if 10 <= len(line) <= 150:
        score += 2
    elif 5 <= len(line) <= 200:
        score += 1

    # 首字母大写
    if line[0].isupper():
        score += 1

    # 不以标点结尾（除问号外）
    if line[-1] not in '.!;,:' or line[-1] == '?':
        score += 1

    # 包含常见标题词汇
    title_keywords = ['analysis', 'study', 'novel', 'approach', 'method',
                      'design', 'optimization', 'performance', 'evaluation',
                      'optical', 'lens', 'system', 'model', 'paper']
    if any(kw in line.lower() for kw in title_keywords):
        score += 1

    # 不含太多数字（标题通常数字较少）
    digit_ratio = sum(c.isdigit() for c in line) / len(line)
    if digit_ratio < 0.1:
        score += 1

    return score


def _extract_title_from_pdf(file_path: str) -> str:
    """从PDF文件中提取标题"""
    try:
        import PyPDF2
    except ImportError:
        print("警告: 未安装PyPDF2，无法从PDF提取标题。请运行: pip install PyPDF2")
        return Path(file_path).stem

    try:
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)

            # 1. 尝试从 PDF 元数据获取标题
            if pdf_reader.metadata and pdf_reader.metadata.title:
                title = pdf_reader.metadata.title.strip()
                if title and len(title) > 5:
                    return title

            # 2. 从第一页文本提取
            first_page = pdf_reader.pages[0]
            text = first_page.extract_text()

            lines = text.split('\n')
            candidates = []

            for line in lines[:15]:  # 扩大到前15行
                line = line.strip()
                if not line:
                    continue

                # 过滤明显不是标题的行
                if _is_likely_not_title(line):
                    continue

                # 计算标题得分
                score = _title_score(line)
                if score > 0:
                    candidates.append((line, score))

            # 返回得分最高的候选
            if candidates:
                candidates.sort(key=lambda x: x[1], reverse=True)
                return candidates[0][0]

            return Path(file_path).stem

    except Exception as e:
        print(f"从PDF文件提取标题失败: {e}")
        return Path(file_path).stem


def _extract_title_from_docx(file_path: str) -> str:
    """从DOCX文件中提取标题"""
    try:
        from docx import Document
    except ImportError:
        print("警告: 未安装python-docx，无法从DOCX提取标题。请运行: pip install python-docx")
        return Path(file_path).stem

    try:
        doc = Document(file_path)

        # 获取第一个段落
        if doc.paragraphs:
            first_paragraph = doc.paragraphs[0].text.strip()
            if first_paragraph and len(first_paragraph) > 5:
                return first_paragraph

        return Path(file_path).stem

    except Exception as e:
        print(f"从DOCX文件提取标题失败: {e}")
        return Path(file_path).stem


def process_papers_directory(input_dir: str, output_file: str) -> None:
    """
    处理论文目录，生成JSON数据集

    Args:
        input_dir: 包含论文的目录路径
        output_file: 输出的JSON文件路径
    """
    input_path = Path(input_dir)

    # 检查输入目录是否存在
    if not input_path.exists():
        raise FileNotFoundError(f"输入目录不存在: {input_dir}")

    if not input_path.is_dir():
        raise NotADirectoryError(f"路径不是目录: {input_dir}")

    papers_data = []

    # 遍历目录中的所有文件
    for file_path in input_path.iterdir():
        if file_path.is_file():
            # 假设所有文件都是论文文件
            title = extract_paper_title(str(file_path))
            location = str(file_path.absolute())

            paper_info = {
                "title": title,
                "location": location
            }
            papers_data.append(paper_info)

    # 创建输出目录（如果不存在）
    output_path = Path(output_file)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # 写入JSON文件
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(papers_data, f, ensure_ascii=False, indent=2)

    print(f"成功处理 {len(papers_data)} 篇论文，数据已保存到: {output_file}")


def main():
    parser = argparse.ArgumentParser(description='将论文文件夹转换为JSON数据集')
    parser.add_argument('--input_dir', type=str, default="dataset/paper_info_extract/data_v1", help='包含论文的目录路径')
    parser.add_argument('--output_file', type=str, default="dataset/paper_info_extract/dataset_json/dataset_v1.json", help='输出的JSON文件路径')

    args = parser.parse_args()

    try:
        process_papers_directory(args.input_dir, args.output_file)
    except Exception as e:
        print(f"错误: {e}")
        return 1

    return 0


if __name__ == "__main__":
    exit(main())
